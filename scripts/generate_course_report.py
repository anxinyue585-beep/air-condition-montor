from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "assets"
OUT_DOCX = REPORT_DIR / "互联网大数据应用技术实践课程报告-城市空气质量全景洞察.docx"

FONT_NAME = "Microsoft YaHei"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
BOLD_FONT_PATH = Path(r"C:\Windows\Fonts\msyhbd.ttc")
if not FONT_PATH.exists():
    FONT_NAME = "Noto Sans SC"
    FONT_PATH = Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf")
    BOLD_FONT_PATH = FONT_PATH

NAVY = RGBColor(15, 23, 42)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 116, 139)
TEAL = RGBColor(15, 118, 110)
ROSE = RGBColor(190, 18, 60)
AMBER = RGBColor(180, 83, 9)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F8FAFC"
MID_GRAY = "E2E8F0"


def load_json(rel: str):
    return json.loads((ROOT / rel).read_text(encoding="utf-8"))


def load_csv(rel: str):
    with (ROOT / rel).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


manifest = load_json("data/processed/dataset_manifest.json")
quality = load_json("data/processed/data_quality_report.json")
processing = load_json("data/processing_results/processing_algorithm_report.json")
analysis = load_json("data/analysis_results/analysis_algorithm_report.json")
risk_rows = load_csv("data/analysis_results/city_risk_ranking.csv")[:10]
cluster_rows = load_csv("data/analysis_results/cluster_summary.csv")
kmeans_rows = load_csv("data/analysis_results/kmeans_parameter_eval.csv")
logistic_rows = load_csv("data/analysis_results/logistic_parameter_eval.csv")
ridge_rows = load_csv("data/analysis_results/ridge_parameter_eval.csv")


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, widths_dxa, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_grid(table, widths_dxa)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if header and r_idx == 0:
                shade_cell(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run_font(run, size=9.2, bold=(header and r_idx == 0), color=NAVY)


def add_text(doc, text, size=10.5, bold=False, color=NAVY, align=None, after=6, before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY)
    return p


def add_caption(doc, text):
    return add_text(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def add_callout(doc, title, body, fill="F0FDFA"):
    table = doc.add_table(rows=1, cols=1)
    format_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=TEAL)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=NAVY)
    add_text(doc, "", after=2)


def add_simple_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    format_table(table, widths_dxa, header=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def draw_rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def pil_font(size=28, bold=False):
    path = BOLD_FONT_PATH if bold and BOLD_FONT_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(path), size)


def draw_text(draw, xy, text, font, fill):
    draw.text(xy, text, font=font, fill=fill)


def create_risk_chart(path: Path):
    width, height = 1300, 760
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(34, True)
    label = pil_font(23)
    small = pil_font(19)
    draw_text(d, (52, 36), "城市污染风险 Top 10", title, "#0f172a")
    draw_text(d, (52, 83), "风险分综合平均 AQI、最高 AQI、污染率和 PM2.5 水平", small, "#64748b")
    max_score = max(float(r["risk_score"]) for r in risk_rows)
    x0, y0, bar_w, gap = 210, 140, 920, 46
    for i, row in enumerate(risk_rows):
        y = y0 + i * 56
        score = float(row["risk_score"])
        city = row["city"]
        color = "#dc2626" if i < 3 else "#f59e0b" if i < 6 else "#14b8a6"
        draw_text(d, (52, y + 9), f"{i + 1}. {city}", label, "#334155")
        draw_rounded_rect(d, (x0, y + 7, x0 + bar_w, y + 36), 14, "#f1f5f9")
        draw_rounded_rect(d, (x0, y + 7, x0 + int(bar_w * score / max_score), y + 36), 14, color)
        draw_text(d, (x0 + bar_w + 24, y + 7), f"{score:.1f}", label, "#0f172a")
    img.save(path)


def create_cluster_chart(path: Path):
    width, height = 1200, 620
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(34, True)
    label = pil_font(23)
    small = pil_font(19)
    colors = {"低污染稳定型": "#10b981", "中等波动型": "#f59e0b", "高污染风险型": "#ef4444"}
    draw_text(d, (52, 36), "K-Means 城市污染画像聚类", title, "#0f172a")
    draw_text(d, (52, 83), "k=3：低污染稳定型、中等波动型、高污染风险型", small, "#64748b")
    max_count = max(int(r["city_count"]) for r in cluster_rows)
    x0, y0, bar_w = 310, 160, 690
    for i, row in enumerate(sorted(cluster_rows, key=lambda r: int(r["city_count"]), reverse=True)):
        y = y0 + i * 115
        name = row["cluster_name"]
        count = int(row["city_count"])
        avg_aqi = float(row["avg_aqi"])
        polluted = float(row["avg_polluted_rate"]) * 100
        color = colors.get(name, "#0ea5e9")
        draw_text(d, (52, y + 8), name, label, "#334155")
        draw_text(d, (52, y + 42), f"均值AQI {avg_aqi:.2f} | 污染率 {polluted:.1f}%", small, "#64748b")
        draw_rounded_rect(d, (x0, y + 14, x0 + bar_w, y + 50), 18, "#f1f5f9")
        draw_rounded_rect(d, (x0, y + 14, x0 + int(bar_w * count / max_count), y + 50), 18, color)
        draw_text(d, (x0 + bar_w + 24, y + 10), f"{count} 城市", label, "#0f172a")
        draw_text(d, (x0, y + 63), f"代表城市：{row['representative_cities']}", small, "#475569")
    img.save(path)


def create_model_chart(path: Path):
    width, height = 1260, 620
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(34, True)
    label = pil_font(24, True)
    value_font = pil_font(36, True)
    small = pil_font(19)
    draw_text(d, (52, 36), "机器学习与预测模型结果", title, "#0f172a")
    draw_text(d, (52, 83), "Logistic Regression 用于下月污染风险分类，Ridge Regression 用于下月 AQI 预测", small, "#64748b")
    cards = [
        ("Logistic Accuracy", f"{analysis['supervised_learning']['logistic_regression']['accuracy']:.4f}", "#0ea5e9"),
        ("Logistic F1", f"{analysis['supervised_learning']['logistic_regression']['f1']:.4f}", "#14b8a6"),
        ("Ridge RMSE", f"{analysis['supervised_learning']['ridge_regression']['rmse']:.4f}", "#f97316"),
        ("Ridge R2", f"{analysis['supervised_learning']['ridge_regression']['r2']:.4f}", "#8b5cf6"),
    ]
    for i, (name, value, color) in enumerate(cards):
        x = 52 + (i % 2) * 590
        y = 150 + (i // 2) * 190
        draw_rounded_rect(d, (x, y, x + 535, y + 145), 22, "#f8fafc", "#e2e8f0", 2)
        draw_rounded_rect(d, (x, y, x + 10, y + 145), 5, color)
        draw_text(d, (x + 34, y + 28), name, label, "#334155")
        draw_text(d, (x + 34, y + 75), value, value_font, color)
    img.save(path)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10.5, color=NAVY)
    normal = styles["Normal"].paragraph_format
    normal.space_after = Pt(6)
    normal.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        set_style_font(styles[name], size=size, color=color, bold=True)
        pf = styles[name].paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.15
    set_style_font(styles["List Bullet"], size=10.5, color=NAVY)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("城市空气质量全景洞察 | 互联网大数据应用技术实践")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("课程报告")
    set_run_font(run, size=9, color=MUTED)
    return doc


def add_cover(doc: Document):
    add_text(doc, "《互联网大数据应用技术实践》", size=13, bold=True, color=TEAL, after=10)
    add_text(doc, "课程报告", size=28, bold=True, color=NAVY, after=4)
    add_text(doc, "城市空气质量大数据分析与可视化系统", size=17, bold=True, color=DARK_BLUE, after=18)
    add_text(
        doc,
        "围绕城市空气质量数据完成数据建设、处理分析与前端展示，并提供 Hadoop/Hive/Spark 平台部署设计和待验收脚本。",
        size=11,
        color=MUTED,
        after=18,
    )
    rows = [
        ("项目名称", "城市空气质量全景洞察"),
        ("课程名称", "互联网大数据应用技术实践"),
        ("团队成员", "刘嘉晨（组长）、顾跃庭（组员）"),
        ("技术路线", "Vue 3 + TypeScript + ECharts + Hadoop + Hive + Spark + 机器学习算法"),
        ("报告日期", "2026年7月11日"),
    ]
    add_simple_table(doc, ["项目", "内容"], rows, [1700, 7660])
    add_callout(
        doc,
        "报告定位",
        "本报告按课程评分细则组织内容，重点说明数据来源、数据规模、平台部署设计与实际验收状态、处理算法、分析模型和可视化系统，并区分已验证结果与待验证项目。",
    )
    doc.add_page_break()


def add_rubric_mapping(doc: Document):
    add_heading(doc, "摘要与评分点覆盖情况", 1)
    add_text(
        doc,
        "本项目以城市空气质量数据为对象，构建了覆盖 60 个城市、120 个监测点、2025 全年小时级记录的百万级实验数据集。"
        "项目已完成数据生成、清洗转换、仓库聚合、Top-N、K-Means、Logistic Regression、Ridge Regression 和 Vue 3 可视化。"
        "Hadoop/Hive/Spark 部分已提供 Docker Compose、HDFS、Hive、Spark 和证据采集脚本，但因当前验收机未安装 Docker，尚未完成容器启动与平台运行验收。",
    )
    rows = [
        ("项目选题与需求分析", "10", "围绕城市空气质量监测、风险识别和可视化决策展开，需求包含查询、筛选、算法分析和多图表展示。"),
        ("数据集建设", "16", "清洗数据 1,051,200 条，提供 Raw/Processed/Warehouse/App Sample 分层存储、字段字典和质量报告。"),
        ("Hadoop/Hive/Spark 平台", "16", "已提供 Docker Compose、HDFS 上传、Hive SQL、Spark ETL 与一键验收脚本；实际平台运行待 Docker 环境补齐后验证。"),
        ("数据处理算法设计", "20", "覆盖清洗、转换、聚合、筛选和优化 5 类方法，含分区索引、Spark cache 与 Parquet 输出。"),
        ("数据分析算法应用", "20", "实现 Top-N、K-Means、Logistic Regression 和 Ridge Regression，包含参数对比和实验指标。"),
        ("数据可视化", "8", "实现数据大屏、明细查询、数据集展示和算法分析页面，包含折线图、柱状图、雷达图、平行坐标等。"),
        ("创新性", "6", "引入机器学习分类、时间序列预测、统一数据源和可复现的数据与算法脚本链路。"),
        ("项目答辩", "4", "已整理运行命令、验证清单和材料目录，可用于录屏答辩展示。"),
    ]
    add_simple_table(doc, ["一级指标", "分值", "当前状态与证据"], rows, [2100, 800, 6460])
    add_text(
        doc,
        "从评分细则看，本项目已满足“至少实现一种数据分析算法”的硬性要求，并进一步实现机器学习分类与预测分析，具备优秀等级所需的算法优化和参数分析材料。",
        bold=True,
        color=TEAL,
    )


def add_project_section(doc: Document):
    add_heading(doc, "一、项目选题与需求分析", 1)
    add_heading(doc, "1.1 选题背景", 2)
    add_text(
        doc,
        "空气质量数据具有明显的时间连续性、空间差异性和多指标相关性，适合用于大数据存储、批处理、统计分析、机器学习和可视化展示。"
        "本项目选择“城市空气质量全景洞察”作为课程大作业主题，目标是通过大数据技术对城市 AQI、PM2.5、PM10、SO2、NO2 等指标进行多维分析，识别高污染城市、污染物结构和未来风险。",
    )
    add_heading(doc, "1.2 功能需求", 2)
    for item in [
        "数据建设需求：构造可复现、字段完整、规模达到课程要求的空气质量实验数据集。",
        "平台需求：使用 Hadoop HDFS 保存百万级数据，使用 Hive 建表查询，使用 Spark 完成分布式 ETL 和聚合输出。",
        "算法需求：实现数据清洗、转换、聚合、筛选和优化；至少实现一种分析算法，并说明原理、参数与结果。",
        "展示需求：通过前端系统提供数据大屏、数据明细、数据集检索、算法分析和项目说明页面。",
        "答辩需求：材料能够被录屏讲解，命令、数据文件、算法结果和页面功能均可被复现。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3 技术路线", 2)
    add_text(
        doc,
        "技术路线采用“数据源 - 数据清洗 - 仓库聚合 - Hadoop/Hive/Spark 平台 - 算法分析 - 前端可视化”的分层方案。"
        "后端侧通过 Python 脚本生成和处理数据，平台侧通过 Docker Compose 管理 Hadoop/Hive/Spark 服务，前端侧使用 Vue 3、TypeScript、Tailwind CSS 和 ECharts 进行交互式展示。",
    )


def add_data_section(doc: Document):
    add_heading(doc, "二、数据集建设", 1)
    summary = manifest["summary"]
    clean_csv = ROOT / "data/processed/air_quality_clean.csv"
    if clean_csv.exists():
        data_size_note = f"CSV 文件约 {clean_csv.stat().st_size / 1024 / 1024:.1f} MB"
    else:
        data_size_note = "主 CSV 未纳入 Git，可通过 build_air_quality_dataset.py 重新生成"
    add_heading(doc, "2.1 数据来源与生成口径", 2)
    add_text(
        doc,
        "项目数据来源由原始种子样本、公开空气质量数据结构参考和 AQI 技术规定共同构成。完整数据由脚本可复现生成，保留空气质量监测数据常见的城市、监测点、时间、污染物和气象字段。"
    )
    rows = [
        ("数据名称", "城市空气质量小时级实验数据集"),
        ("时间范围", f"{summary['date_range'][0]} 至 {summary['date_range'][1]}"),
        ("时间粒度", "小时级"),
        ("空间范围", f"{summary['city_count']} 个城市、{summary['station_count']} 个监测点"),
        ("主数据规模", f"{summary['clean_rows']:,} 条，{data_size_note}"),
        ("聚合结果", f"城市日表 {summary['city_day_rows']:,} 行，城市月表 {summary['city_month_rows']:,} 行"),
    ]
    add_simple_table(doc, ["项目", "说明"], rows, [1800, 7560])

    add_heading(doc, "2.2 数据清洗与质量控制", 2)
    q = quality["cleaning_actions"]
    checks = quality["checks"]
    rows = [
        ("去重", f"演示样本删除重复记录 {q['duplicates_removed']} 行；完整数据重复记录数为 0。"),
        ("缺失值处理", f"演示样本填充缺失值 {q['missing_values_filled']} 个，采用城市中位数或全局中位数。"),
        ("异常值处理", f"负数修正 {q['negative_values_fixed']} 个，AQI 异常截断 {q['outliers_capped']} 行。"),
        ("等级校验", "根据 AQI 重算空气质量等级，完整数据等级不一致记录数为 0。"),
        ("范围检查", f"AQI 范围为 {checks['aqi_range']}，清洗后无缺失值。"),
    ]
    add_simple_table(doc, ["质量控制项", "处理结果"], rows, [1900, 7460])

    add_heading(doc, "2.3 存储结构", 2)
    rows = [
        ("Raw", "data/raw/", "保留原始种子数据、脏样本和来源清单。"),
        ("Processed", "data/processed/", "保留清洗后的完整小时级数据、前端抽样数据和质量报告。"),
        ("Warehouse", "data/warehouse/", "保留城市日表、城市月表、区域月表、Top-N 和分区索引。"),
        ("Docs", "docs/", "保留数据来源、字段字典、数据建设报告和平台报告。"),
    ]
    add_simple_table(doc, ["层级", "路径", "说明"], rows, [1200, 2200, 5960])
    add_callout(
        doc,
        "数据建设结论",
        "数据规模超过课程建议的 CSV 100 万条要求；清洗流程覆盖去重、缺失值、异常值和字段转换；存储结构具备原始、清洗、仓库和前端抽样四层，便于追溯和展示。",
    )


def add_platform_section(doc: Document):
    add_heading(doc, "三、Hadoop/Hive/Spark 平台部署", 1)
    add_text(
        doc,
        "平台设计采用 Docker Compose，计划集成 Hadoop HDFS、YARN、Hive Metastore、HiveServer2、Spark Master 和 Spark Worker。"
        "配置目标是存储百万级清洗数据、执行 Hive SQL 查询和 Spark 分布式 ETL；当前仅完成配置、脚本与数据前置检查，尚未完成容器运行验收。",
    )
    rows = [
        ("Hadoop NameNode", "aq-namenode", "HDFS 元数据管理，端口 9870/9000"),
        ("Hadoop DataNode", "aq-datanode", "HDFS 数据块存储，端口 9864"),
        ("YARN ResourceManager", "aq-resourcemanager", "资源调度，端口 8088"),
        ("Hive Metastore + DB", "aq-hive-metastore / aq-hive-postgres", "Hive 元数据服务与 PostgreSQL 元数据库"),
        ("HiveServer2", "aq-hive-server", "Hive SQL 服务，端口 10000"),
        ("Spark Master/Worker", "aq-spark-master / aq-spark-worker", "Spark 作业调度与执行，端口 8080/8081"),
    ]
    add_simple_table(doc, ["服务", "容器名", "作用"], rows, [2100, 3000, 4260])
    add_heading(doc, "3.1 HDFS 与 Hive 设计", 2)
    add_text(
        doc,
        "上传脚本设计为将 `air_quality_clean.csv` 写入 HDFS 路径 `/warehouse/air_quality/clean/air_quality_clean.csv`。Hive SQL 设计为创建外部表 `air_quality_clean_csv`、城市月度聚合表 `air_quality_city_month` 与城市平均 AQI Top-N 表 `air_quality_city_topn`；实际行数查询待 Docker 平台运行后验收。"
    )
    add_heading(doc, "3.2 Spark ETL 设计", 2)
    add_text(
        doc,
        "Spark 作业 `scripts/spark_air_quality_etl.py` 设计为读取 HDFS CSV，执行字段类型转换、清洗、缓存、分区、城市日/月聚合和 Top-N 输出。Parquet 与 CSV 输出路径已定义，但尚无真实容器执行记录。"
    )
    rows = [
        ("启动平台", "powershell -ExecutionPolicy Bypass -File scripts\\platform_up.ps1"),
        ("上传 HDFS", "powershell -ExecutionPolicy Bypass -File scripts\\upload_to_hdfs.ps1"),
        ("Hive 建表查询", "powershell -ExecutionPolicy Bypass -File scripts\\run_hive_air_quality.ps1"),
        ("Spark 作业", "powershell -ExecutionPolicy Bypass -File scripts\\run_spark_air_quality.ps1"),
        ("冒烟检查", "powershell -ExecutionPolicy Bypass -File scripts\\platform_smoke_check.ps1"),
    ]
    add_simple_table(doc, ["步骤", "命令"], rows, [1800, 7560])
    add_callout(
        doc,
        "平台当前状态",
        "项目已提供平台配置、启动脚本、HDFS 上传脚本、Hive SQL、Spark ETL 作业和一键证据采集命令。2026-07-11 数据前置检查已通过，但 Docker CLI 缺失，HDFS、Hive、Spark 均未实际启动，因此本报告不将平台部署表述为已完成。",
        fill="F8FAFC",
    )


def add_processing_section(doc: Document):
    add_heading(doc, "四、数据处理算法设计", 1)
    add_text(
        doc,
        "数据处理算法部分对应评分细则中的清洗、转换、聚合、筛选和优化 5 个评分项。项目不仅满足至少 3 种方法的最低要求，而且完整覆盖全部 5 类处理方法。"
    )
    rows = []
    for item in processing["rubric_mapping"]:
        rows.append((item["item"], item["score"], "、".join(item["methods"]), "；".join(item["evidence"])))
    add_simple_table(doc, ["评分项", "分值", "已实现方法", "证据文件"], rows, [1200, 700, 3200, 4260])

    add_heading(doc, "4.1 清洗与转换", 2)
    add_text(
        doc,
        "清洗算法按记录编号和业务字段去重，使用城市中位数或全局中位数填充缺失值，对负数和超范围 AQI 进行修正，并根据 AQI 重新计算等级。"
        "转换算法输出 Min-Max 归一化、Z-Score 标准化、等级编码、早晚高峰标记和重污染标记等特征。"
    )
    trans = processing["transformation"]
    rows = [
        ("特征样本行数", f"{trans['feature_sample_rows']:,}"),
        ("等级编码", "优=0，良=1，污染=2"),
        ("核心数值特征", "AQI、PM2.5、PM10"),
        ("输出文件", trans["output"]),
    ]
    add_simple_table(doc, ["项目", "结果"], rows, [2100, 7260])

    add_heading(doc, "4.2 聚合、筛选与优化", 2)
    agg = processing["aggregation"]
    filt = processing["filtering"]
    opt = processing["optimization"]
    rows = [
        ("城市月度聚合", f"{agg['city_month_rows']} 行"),
        ("区域月度聚合", f"{agg['region_month_rows']} 行"),
        ("城市 AQI Top-N", f"{agg['city_topn_rows']} 行"),
        ("高污染样本", f"{filt['high_pollution_rows']:,} 条，规则：{filt['rule']}"),
        ("早晚高峰高污染样本", f"{filt['rush_hour_high_pollution_rows']:,} 条"),
        ("分区优化", f"{' + '.join(opt['partition_keys'])}，共 {opt['partition_count']} 个逻辑分区"),
        ("Spark 优化", f"shuffle partitions={opt['shuffle_partitions_for_spark']}，使用 cache() 与 repartition()"),
    ]
    add_simple_table(doc, ["处理项", "结果"], rows, [2300, 7060])


def add_analysis_section(doc: Document):
    add_heading(doc, "五、数据分析算法应用", 1)
    ds = analysis["supervised_learning"]["dataset"]
    add_text(
        doc,
        f"分析算法输入为城市月度聚合表，共 720 条城市月度记录。监督学习任务采用“当前月特征预测下月结果”的样本构造方式，样本数 {ds['sample_count']}，训练集 {ds['train_count']}，验证集 {ds['validation_count']}，测试集 {ds['test_count']}。数据按目标月份顺序划分，超参数仅在验证集选择，测试集只用于最终评价。"
    )
    add_simple_table(
        doc,
        ["算法类别", "算法", "用途", "证据文件"],
        [
            ("A 类基础算法", "Top-N 风险排名", "识别污染风险最高城市", "city_risk_ranking.csv"),
            ("B 类数据挖掘算法", "K-Means 聚类", "划分城市污染画像类型", "kmeans_parameter_eval.csv / cluster_summary.csv"),
            ("C 类机器学习算法", "Logistic Regression", "预测下月污染风险", "logistic_parameter_eval.csv"),
            ("预测创新项", "Ridge Regression", "预测下月平均 AQI", "ridge_parameter_eval.csv / predictions.csv"),
        ],
        [1800, 2200, 2500, 2860],
    )

    add_heading(doc, "5.1 Top-N 城市污染风险排名", 2)
    add_text(
        doc,
        "Top-N 排名使用平均 AQI、最高 AQI、污染率和 PM2.5 均值构建综合风险分。归一化后按权重求和，最终识别出石家庄、太原、乌鲁木齐、西安、郑州等重点风险城市。"
    )
    img = ASSET_DIR / "risk_top10.png"
    doc.add_picture(str(img), width=Inches(6.35))
    add_caption(doc, "图 1 城市污染风险 Top 10")
    add_simple_table(
        doc,
        ["排名", "城市", "区域", "风险分", "平均AQI", "污染率"],
        [(r["rank"], r["city"], r["region"], r["risk_score"], r["avg_aqi"], f"{float(r['polluted_rate']) * 100:.1f}%") for r in risk_rows[:5]],
        [800, 1400, 1200, 1400, 1400, 1600],
    )

    add_heading(doc, "5.2 K-Means 城市污染画像聚类", 2)
    add_text(
        doc,
        "K-Means 使用平均 AQI、最高 AQI、污染率、PM2.5、PM10、SO2、NO2 等特征，先进行 Z-Score 标准化，再测试 k=2 至 k=6。项目选择 k=3，因为三类结果具有较强业务解释性。"
    )
    doc.add_picture(str(ASSET_DIR / "cluster_summary.png"), width=Inches(6.25))
    add_caption(doc, "图 2 K-Means 城市污染画像聚类结果")
    add_simple_table(
        doc,
        ["k", "SSE", "轮廓系数", "迭代次数"],
        [(r["k"], r["inertia_sse"], r["silhouette"], r["iterations"]) for r in kmeans_rows],
        [1000, 2600, 2600, 1600],
    )

    add_heading(doc, "5.3 Logistic Regression 下月污染风险分类", 2)
    logi = analysis["supervised_learning"]["logistic_regression"]
    add_text(
        doc,
        f"Logistic Regression 将 next_month_avg_aqi >= 100 定义为下月污染风险。验证集选择 lambda={logi['best_lambda']}（验证 F1={logi['validation_f1']:.4f}），训练集与验证集重训后，隔离测试集 Accuracy={logi['accuracy']:.4f}，Precision={logi['precision']:.4f}，Recall={logi['recall']:.4f}，F1={logi['f1']:.4f}。"
        f"验证期正样本仅 {ds['validation_positive_count']} 条，参数区分度有限；实验结果说明模型误报率较低，但仍有部分污染月份未被提前捕捉，后续应引入多年真实数据并采用滚动时间验证。下表为验证集参数评估结果。"
    )
    add_simple_table(
        doc,
        ["lambda", "Accuracy", "Precision", "Recall", "F1"],
        [(r["lambda"], r["accuracy"], r["precision"], r["recall"], r["f1"]) for r in logistic_rows],
        [1200, 2000, 2000, 2000, 1600],
    )

    add_heading(doc, "5.4 Ridge Regression 下月 AQI 预测", 2)
    ridge = analysis["supervised_learning"]["ridge_regression"]
    add_text(
        doc,
        f"Ridge Regression 用于预测下月平均 AQI，通过 L2 正则化缓解多污染物特征之间的相关性。验证集选择 alpha={ridge['best_alpha']}（验证 RMSE={ridge['validation_rmse']:.4f}），训练集与验证集重训后，隔离测试集 MAE={ridge['mae']:.4f}，RMSE={ridge['rmse']:.4f}，R2={ridge['r2']:.4f}。"
        f"与直接使用当前月 AQI 的基线相比，RMSE 从 {ridge['baseline_rmse']:.4f} 降低到 {ridge['rmse']:.4f}。下表为验证集参数评估结果。"
    )
    doc.add_picture(str(ASSET_DIR / "model_metrics.png"), width=Inches(6.35))
    add_caption(doc, "图 3 机器学习与预测模型指标")
    add_simple_table(
        doc,
        ["alpha", "MAE", "RMSE", "R2"],
        [(r["alpha"], r["mae"], r["rmse"], r["r2"]) for r in ridge_rows],
        [1200, 2400, 2400, 1800],
    )


def add_visualization_section(doc: Document):
    add_heading(doc, "六、数据可视化系统展示", 1)
    add_text(
        doc,
        "前端系统使用 Vue 3、TypeScript、Vite、Tailwind CSS 与 ECharts 6 实现，页面围绕课程评分中的数据展示、查询、筛选、分页、导出和算法解释需求设计。"
    )
    rows = [
        ("数据大屏", "/dashboard", "展示 AQI 趋势、城市质量结构、污染物占比、雷达图和平行坐标画像，支持周/月/全年维度切换。"),
        ("数据明细", "/data", "基于统一数据源进行城市和等级筛选，提供分页浏览和 CSV 导出。"),
        ("数据集展示", "/dataset", "支持关键字、城市、等级、季节筛选，提供排序、分页、JSON/CSV 导出和原始数据抽屉。"),
        ("算法分析", "/analysis", "集中展示 Top-N、K-Means、Logistic Regression 和 Ridge Regression 的参数与结果。"),
        ("关于项目", "/about", "说明技术路线、模块职责、算法成果和团队分工。"),
    ]
    add_simple_table(doc, ["页面", "路由", "功能说明"], rows, [1500, 1400, 6460])
    add_heading(doc, "6.1 可视化设计说明", 2)
    for item in [
        "AQI 趋势概览：展示整体均值、高风险城市和低风险城市对比，并标注 AQI 100 风险线。",
        "城市空气质量结构：用绿色和红色区分优良与污染记录，全年维度自动切换为月份口径。",
        "污染物雷达图：比较 AQI、PM2.5、PM10、NO2、SO2 的综合特征。",
        "重点城市污染画像：采用平行坐标展示多指标差异，并增加固定城市标签，减少悬停依赖。",
        "算法分析中心：以风险条、参数表、聚类摘要和预测结果表解释模型输出。",
    ]:
        add_bullet(doc, item)


def add_innovation_section(doc: Document):
    add_heading(doc, "七、创新性与项目特色", 1)
    rows = [
        ("可扩展大数据链路", "数据生成、清洗、仓库和前端展示已形成可运行链路；Hadoop/Hive/Spark 接口与脚本已预留，平台闭环待 Docker 验收。"),
        ("机器学习算法", "实现 Logistic Regression 下月污染风险分类，满足优秀等级中的机器学习算法要求。"),
        ("预测分析", "实现 Ridge Regression 下月 AQI 预测，属于时间序列预测方向的创新扩展。"),
        ("参数对比", "K-Means、Logistic Regression、Ridge Regression 均提供参数评估结果。"),
        ("统一数据源", "数据大屏、数据明细、数据集展示和算法分析均使用同一套处理结果，减少前后数据口径不一致。"),
        ("可复现材料", "数据与算法结果可由 Python 脚本复现；平台部分提供 PowerShell、Hive SQL、Spark 作业和验收脚本，待具备 Docker 后复验。"),
    ]
    add_simple_table(doc, ["特色", "说明"], rows, [2100, 7260])


def add_conclusion_section(doc: Document):
    add_heading(doc, "八、总结与展望", 1)
    add_text(
        doc,
        "本项目已完成城市空气质量数据生成、处理分析和可视化系统实现。数据层面已构建百万级小时数据和多级仓库表；算法层面已完成处理算法、数据挖掘、机器学习分类和预测分析；展示层面已完成可交互的数据大屏、明细查询、实时数据和算法分析页面。"
        "平台层面已完成 Hadoop/Hive/Spark 拓扑、Compose 配置、执行脚本和验收脚本，但由于当前机器缺少 Docker，容器启动、HDFS 上传、Hive 查询和 Spark ETL 尚未实际验收。",
    )
    add_text(
        doc,
        "后续可继续优化三个方向：第一，引入真实在线空气质量接口或更多城市历史公开数据，提高数据真实性；第二，加入 Spark Streaming 或 Flink 完成实时数据处理；第三，引入随机森林、XGBoost 或 LSTM 等模型，与当前线性模型进行更深入对比。",
    )


def add_appendix(doc: Document):
    add_heading(doc, "附录：项目材料与运行命令", 1)
    add_heading(doc, "A.1 核心材料清单", 2)
    rows = [
        ("数据建设", "data/raw/、data/processed/、data/warehouse/、docs/data_construction_report.md"),
        ("平台部署", "platform/docker-compose.yml、scripts/platform_up.ps1、sql/hive_air_quality.sql、scripts/spark_air_quality_etl.py"),
        ("数据处理算法", "scripts/run_data_processing_algorithms.py、sql/data_processing_algorithms.sql、data/processing_results/"),
        ("数据分析算法", "scripts/run_data_analysis_algorithms.py、data/analysis_results/、docs/analysis_algorithm_application_report.md"),
        ("前端系统", "src/views/、src/components/、src/composables/、package.json"),
        ("课程报告", str(OUT_DOCX.relative_to(ROOT))),
    ]
    add_simple_table(doc, ["模块", "文件或目录"], rows, [1800, 7560])
    add_heading(doc, "A.2 一键运行命令", 2)
    rows = [
        ("生成数据集", "python scripts\\build_air_quality_dataset.py"),
        ("启动大数据平台", "powershell -ExecutionPolicy Bypass -File scripts\\platform_up.ps1"),
        ("上传到 HDFS", "powershell -ExecutionPolicy Bypass -File scripts\\upload_to_hdfs.ps1"),
        ("运行 Hive", "powershell -ExecutionPolicy Bypass -File scripts\\run_hive_air_quality.ps1"),
        ("运行 Spark", "powershell -ExecutionPolicy Bypass -File scripts\\run_spark_air_quality.ps1"),
        ("运行处理算法", "powershell -ExecutionPolicy Bypass -File scripts\\run_data_processing_algorithms.ps1"),
        ("运行分析算法", "powershell -ExecutionPolicy Bypass -File scripts\\run_data_analysis_algorithms.ps1"),
        ("启动前端", "npm run dev"),
    ]
    add_simple_table(doc, ["任务", "命令"], rows, [1800, 7560])


def create_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_risk_chart(ASSET_DIR / "risk_top10.png")
    create_cluster_chart(ASSET_DIR / "cluster_summary.png")
    create_model_chart(ASSET_DIR / "model_metrics.png")


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    create_assets()
    doc = setup_document()
    add_cover(doc)
    add_rubric_mapping(doc)
    add_project_section(doc)
    add_data_section(doc)
    add_platform_section(doc)
    add_processing_section(doc)
    add_analysis_section(doc)
    add_visualization_section(doc)
    add_innovation_section(doc)
    add_conclusion_section(doc)
    add_appendix(doc)
    doc.core_properties.title = "互联网大数据应用技术实践课程报告"
    doc.core_properties.subject = "城市空气质量大数据分析与可视化系统"
    doc.core_properties.author = "刘嘉晨、顾跃庭"
    doc.core_properties.keywords = "大数据,Hadoop,Hive,Spark,空气质量,机器学习,可视化"
    doc.core_properties.created = datetime.now()
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
